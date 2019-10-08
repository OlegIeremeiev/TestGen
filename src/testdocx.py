import string, os
from docx import Document
from docx.shared import *


class TestDocx:

    def __init__(self):
        self.tests = Document()
        self.answers = Document()

    def write_test(self, data, count):
        table_columns = 10
        folders = ('generated', os.path.sep + data['general']['course'])

        self.__set_margins(self.tests, [1, 1, 1, 1])
        self.__set_margins(self.answers, [1, 1, 1, 1])
        self.__set_style_normal(self.tests)
        self.__set_style_normal(self.answers)
        alphabet = string.ascii_uppercase

        for variant in range(count):

            # title section
            tests_header = data['general']['discipline'] + ", " + data['general']['title']
            tests_header += '\t\tВариант №' + str(variant + 1)
            self.__set_paragraph(self.tests.add_paragraph(tests_header), (1, 0, 12))

            text = 'ФИО:\t\t\t\t\tГруппа:\t\t\tДата:'
            self.__set_paragraph(self.tests.add_paragraph(text), (1, 0, 12))

            self.answers.add_paragraph(tests_header)
            self.answers.add_paragraph()
            # table section
            columns = min(len(data['questions'][0]), table_columns)
            table = self.tests.add_table(0, columns + 1)
            table.style = 'Table Grid'
            self.tests.add_paragraph()

            table_answers = self.answers.add_table(0, columns + 1)
            table_answers.style = 'Table Grid'
            self.answers.add_paragraph()
            self.answers.add_paragraph()

            # questions section
            questions = data['questions'][variant % len(data['questions'])]
            table_data = [[]] * len(questions)
            for index, question in zip(list(range(len(questions))), questions):
                self.tests.add_paragraph(str(index + 1) + '.\t' + question.question) #+ " / " + str(question.index))
                for ind, ans in zip(alphabet, question.answers):
                    txt = ind + ")\t" + ans['text'] # + " / " + str(ans['right'])
                    self.tests.add_paragraph(txt)
                self.tests.add_paragraph()
                table_data[index] = [index, question.cost, question.correct]

            # table filling section
            row_begin = -3
            col_ind = 0
            for ind, cost, corr in table_data:
                if (ind % columns) == 0:
                    table.add_row()
                    table.add_row()
                    table.add_row()
                    table_answers.add_row()
                    table_answers.add_row()
                    table_answers.add_row()
                    row_begin += 3
                    self.__set_cell_paragraph(table.cell(row_begin, 0), '№', space=(1, 3, 3))
                    self.__set_cell_paragraph(table.cell(row_begin + 1, 0), 'Ответ', space=(1, 12, 12))
                    self.__set_cell_paragraph(table.cell(row_begin + 2, 0), 'Баллы', space=(1, 3, 3))
                    col_ind = 1
                    self.__set_cell_paragraph(table_answers.cell(row_begin, 0), '№', space=(1, 1, 1))
                    self.__set_cell_paragraph(table_answers.cell(row_begin + 1, 0), 'Ответ', space=(1, 1, 1))
                    self.__set_cell_paragraph(table_answers.cell(row_begin + 2, 0), 'Баллы', space=(1, 1, 1))

                self.__set_cell_paragraph(table.cell(row_begin, col_ind), str(ind + 1))
                self.__set_cell_paragraph(table.cell(row_begin + 2, col_ind), str(cost))

                self.__set_cell_paragraph(table_answers.cell(row_begin, col_ind), str(ind + 1))
                self.__set_cell_paragraph(table_answers.cell(row_begin + 1, col_ind), corr)#chr(ord('A') + corr))
                self.__set_cell_paragraph(table_answers.cell(row_begin + 2, col_ind), str(cost))
                col_ind += 1

            if variant < count - 1:
                self.tests.add_page_break()

        tmp=''
        for fold in folders:
            tmp +=fold
            if not os.path.exists(tmp):
                os.makedirs(tmp)

        tmp += os.path.sep + data['general']['filename']
        self.tests.save(tmp + '_tests.docx')
        self.answers.save(tmp + '_answers.docx')

    @staticmethod
    def __set_tab(paragraph, size_tab_cm):
        pt_format = paragraph.paragraph_format
        pt_format.tab_stops.add_tab_stop(Cm(size_tab_cm))

    @staticmethod
    def __set_margins(doc, margins_list_cm):
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(margins_list_cm[0])
            section.bottom_margin = Cm(margins_list_cm[1])
            section.left_margin = Cm(margins_list_cm[2])
            section.right_margin = Cm(margins_list_cm[3])

    @staticmethod
    def __set_paragraph(paragraph, space_list_pt=(1, 0, 0)):
        pt_format = paragraph.paragraph_format
        pt_format.line_spacing = space_list_pt[0]
        pt_format.space_before = Pt(space_list_pt[1])
        pt_format.space_after = Pt(space_list_pt[2])

    @staticmethod
    def __set_font(run, name, size_pt):
        font = run.font
        font.name = name
        font.size = Pt(size_pt)

    def __add_default_paragraph(self, doc, text, space=(1, 0, 0), tab=1, font_name='Times New Roman', font_size=12):
        pt = doc.add_paragraph()
        self.__set_paragraph(pt, space)
        self.__set_tab(pt, tab)
        self.__add_default_run(pt, text, font_name, font_size)

    def __add_default_run(self, paragraph, text, name='Times New Roman', size=12):
        run = paragraph.add_run(text)
        self.__set_font(run, name, size)

    @staticmethod
    def __set_style_normal(doc):
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        pt_format = style.paragraph_format
        pt_format.line_spacing = 1
        pt_format.space_before = Pt(0)
        pt_format.space_after = Pt(0)

    @staticmethod
    def __set_cell_paragraph(cell, text, alignment=1, space=(1, 0, 0)):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = alignment
        TestDocx.__set_paragraph(p, space)
